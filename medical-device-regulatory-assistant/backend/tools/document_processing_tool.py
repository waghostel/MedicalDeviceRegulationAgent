"""
Document Processing Tool for Medical Device Regulatory Assistant

This tool provides comprehensive document processing capabilities including:
- PDF/DOCX to markdown conversion
- OCR functionality for scanned documents
- NLP-based text extraction and structured data parsing
- Document search and relevance scoring
- Document summarization
- Version tracking and change detection
"""

import os
import re
import hashlib
import logging
from typing import Dict, List, Optional, Any, Tuple
from pathlib import Path
from datetime import datetime
import asyncio
from dataclasses import dataclass, asdict

# Document processing imports
from pypdf import PdfReader
from docx import Document as DocxDocument
import pytesseract
from PIL import Image

# NLP and ML imports
import spacy
import nltk
from transformers import pipeline, AutoTokenizer, AutoModel
from sentence_transformers import SentenceTransformer
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np

# Web scraping and caching
import requests
import requests_cache
from bs4 import BeautifulSoup

# LangChain integration
from langchain.tools import BaseTool
from langchain.schema import Document

# Internal imports
try:
    from backend.models.document_models import (
        ProcessedDocument,
        DocumentMetadata,
        DocumentVersion,
        ExtractionResult,
        SummaryResult,
        SearchResult
    )
except ImportError:
    # Fallback for direct script execution
    from models.document_models import (
        ProcessedDocument,
        DocumentMetadata,
        DocumentVersion,
        ExtractionResult,
        SummaryResult,
        SearchResult
    )

# Configure logging
logger = logging.getLogger(__name__)

# Initialize NLP models (lazy loading)
_nlp_model = None
_summarizer = None
_sentence_transformer = None

def get_nlp_model():
    """Lazy load spaCy model"""
    global _nlp_model
    if _nlp_model is None:
        try:
            _nlp_model = spacy.load("en_core_web_sm")
        except OSError:
            logger.warning("spaCy model 'en_core_web_sm' not found. Using blank model.")
            _nlp_model = spacy.blank("en")
    return _nlp_model

def get_summarizer():
    """Lazy load summarization model"""
    global _summarizer
    if _summarizer is None:
        _summarizer = pipeline(
            "summarization",
            model="facebook/bart-large-cnn",
            tokenizer="facebook/bart-large-cnn"
        )
    return _summarizer

def get_sentence_transformer():
    """Lazy load sentence transformer model"""
    global _sentence_transformer
    if _sentence_transformer is None:
        _sentence_transformer = SentenceTransformer('all-MiniLM-L6-v2')
    return _sentence_transformer


@dataclass
class DocumentProcessingConfig:
    """Configuration for document processing"""
    max_file_size_mb: int = 50
    supported_formats: List[str] = None
    ocr_enabled: bool = True
    nlp_enabled: bool = True
    cache_enabled: bool = True
    cache_duration_hours: int = 24
    
    def __post_init__(self):
        if self.supported_formats is None:
            self.supported_formats = ['.pdf', '.docx', '.txt', '.html', '.md']


class DocumentProcessingTool(BaseTool):
    """
    Comprehensive document processing tool for FDA regulatory documents
    """
    
    name: str = "document_processing_tool"
    description: str = """
    Process regulatory documents including PDF/DOCX conversion to markdown,
    OCR for scanned documents, NLP extraction, search, and summarization.
    """
    
    def __init__(self, config: Optional[DocumentProcessingConfig] = None, **kwargs):
        super().__init__(**kwargs)
        # Use object.__setattr__ to bypass Pydantic validation for custom attributes
        object.__setattr__(self, 'config', config or DocumentProcessingConfig())
        
        # Initialize caching
        if self.config.cache_enabled:
            requests_cache.install_cache(
                'document_cache',
                expire_after=self.config.cache_duration_hours * 3600
            )
        
        # Initialize NLTK data
        self._ensure_nltk_data()
    
    def _ensure_nltk_data(self):
        """Ensure required NLTK data is downloaded"""
        try:
            nltk.data.find('tokenizers/punkt')
            nltk.data.find('corpora/stopwords')
        except LookupError:
            nltk.download('punkt', quiet=True)
            nltk.download('stopwords', quiet=True)
    
    def _run(self, action: str, **kwargs) -> Dict[str, Any]:
        """
        Execute document processing action
        
        Args:
            action: The action to perform (convert, extract, search, summarize, etc.)
            **kwargs: Action-specific parameters
        
        Returns:
            Dictionary containing processing results
        """
        try:
            if action == "convert_to_markdown":
                return self.convert_to_markdown(
                    file_path=kwargs.get('file_path'),
                    url=kwargs.get('url'),
                    content=kwargs.get('content')
                )
            elif action == "extract_structured_data":
                return self.extract_structured_data(
                    text=kwargs['text'],
                    extraction_type=kwargs.get('extraction_type', 'regulatory')
                )
            elif action == "search_documents":
                return self.search_documents(
                    query=kwargs['query'],
                    documents=kwargs.get('documents', []),
                    top_k=kwargs.get('top_k', 5)
                )
            elif action == "summarize_document":
                return self.summarize_document(
                    text=kwargs['text'],
                    max_length=kwargs.get('max_length', 150),
                    min_length=kwargs.get('min_length', 50)
                )
            elif action == "track_version":
                return self.track_document_version(
                    document_id=kwargs['document_id'],
                    content=kwargs['content'],
                    metadata=kwargs.get('metadata', {})
                )
            else:
                raise ValueError(f"Unknown action: {action}")
                
        except Exception as e:
            logger.error(f"Error in document processing: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "action": action
            }
    
    async def _arun(self, action: str, **kwargs) -> Dict[str, Any]:
        """Async version of _run"""
        return await asyncio.get_event_loop().run_in_executor(
            None, self._run, action, **kwargs
        )
    
    def convert_to_markdown(
        self,
        file_path: Optional[str] = None,
        url: Optional[str] = None,
        content: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Convert various document formats to markdown
        
        Args:
            file_path: Path to local file
            url: URL to remote document
            content: Raw text content
        
        Returns:
            Dictionary with markdown content and metadata
        """
        try:
            if file_path:
                return self._convert_file_to_markdown(file_path)
            elif url:
                return self._convert_url_to_markdown(url)
            elif content:
                return self._convert_text_to_markdown(content)
            else:
                raise ValueError("Must provide file_path, url, or content")
                
        except Exception as e:
            logger.error(f"Error converting to markdown: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "markdown": "",
                "metadata": {}
            }
    
    def _convert_file_to_markdown(self, file_path: str) -> Dict[str, Any]:
        """Convert local file to markdown"""
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Check file size
        file_size_mb = path.stat().st_size / (1024 * 1024)
        if file_size_mb > self.config.max_file_size_mb:
            raise ValueError(f"File too large: {file_size_mb:.1f}MB > {self.config.max_file_size_mb}MB")
        
        # Check file format
        if path.suffix.lower() not in self.config.supported_formats:
            raise ValueError(f"Unsupported format: {path.suffix}")
        
        # Extract text based on file type
        if path.suffix.lower() == '.pdf':
            text = self._extract_pdf_text(file_path)
        elif path.suffix.lower() == '.docx':
            text = self._extract_docx_text(file_path)
        elif path.suffix.lower() in ['.txt', '.md']:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        else:
            raise ValueError(f"Conversion not implemented for {path.suffix}")
        
        # Convert to markdown
        markdown = self._text_to_markdown(text)
        
        # Generate metadata
        metadata = DocumentMetadata(
            filename=path.name,
            file_size=int(path.stat().st_size),
            file_type=path.suffix.lower(),
            created_at=datetime.fromtimestamp(path.stat().st_ctime),
            modified_at=datetime.fromtimestamp(path.stat().st_mtime),
            content_hash=self._calculate_hash(text)
        )
        
        return {
            "success": True,
            "markdown": markdown,
            "metadata": asdict(metadata),
            "original_text": text
        }
    
    def _convert_url_to_markdown(self, url: str) -> Dict[str, Any]:
        """Convert web document to markdown"""
        try:
            response = requests.get(url, timeout=30)
            response.raise_for_status()
            
            # Parse HTML content
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Extract main content (remove navigation, ads, etc.)
            text = self._extract_main_content(soup)
            
            # Convert to markdown
            markdown = self._text_to_markdown(text)
            
            # Generate metadata
            metadata = DocumentMetadata(
                filename=url.split('/')[-1] or 'web_document',
                file_size=len(response.content),
                file_type='html',
                source_url=url,
                content_hash=self._calculate_hash(text)
            )
            
            return {
                "success": True,
                "markdown": markdown,
                "metadata": asdict(metadata),
                "original_text": text
            }
            
        except Exception as e:
            logger.error(f"Error converting URL to markdown: {str(e)}")
            raise
    
    def _convert_text_to_markdown(self, content: str) -> Dict[str, Any]:
        """Convert raw text to markdown"""
        markdown = self._text_to_markdown(content)
        
        metadata = DocumentMetadata(
            filename="text_content",
            file_size=len(content.encode('utf-8')),
            file_type='text',
            content_hash=self._calculate_hash(content)
        )
        
        return {
            "success": True,
            "markdown": markdown,
            "metadata": asdict(metadata),
            "original_text": content
        }
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    
                    # If text extraction fails, try OCR
                    if not page_text.strip() and self.config.ocr_enabled:
                        logger.info(f"No text found on page {page_num + 1}, attempting OCR")
                        # Note: OCR implementation would require additional setup
                        # This is a placeholder for OCR functionality
                        page_text = self._ocr_pdf_page(file_path, page_num)
                    
                    text += page_text + "\n\n"
                    
        except Exception as e:
            logger.error(f"Error extracting PDF text: {str(e)}")
            raise
        
        return text.strip()
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = DocxDocument(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {str(e)}")
            raise
    
    def _ocr_pdf_page(self, file_path: str, page_num: int) -> str:
        """
        Perform OCR on a PDF page (placeholder implementation)
        
        Note: This requires additional setup including:
        - Tesseract OCR installation
        - PDF to image conversion
        - Image preprocessing
        """
        # This is a simplified placeholder
        # Full implementation would require pdf2image and proper OCR setup
        logger.warning("OCR functionality not fully implemented")
        return ""
    
    def _extract_main_content(self, soup: BeautifulSoup) -> str:
        """Extract main content from HTML, removing navigation and ads"""
        # Remove unwanted elements
        for element in soup(['nav', 'header', 'footer', 'aside', 'script', 'style']):
            element.decompose()
        
        # Try to find main content area
        main_content = soup.find('main') or soup.find('article') or soup.find('div', class_=re.compile(r'content|main|body'))
        
        if main_content:
            return main_content.get_text(separator='\n', strip=True)
        else:
            return soup.get_text(separator='\n', strip=True)
    
    def _text_to_markdown(self, text: str) -> str:
        """Convert plain text to markdown with basic formatting"""
        # Clean up text
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)  # Remove excessive line breaks
        text = re.sub(r'[ \t]+', ' ', text)  # Normalize whitespace
        
        lines = text.split('\n')
        markdown_lines = []
        
        for line in lines:
            line = line.strip()
            if not line:
                markdown_lines.append('')
                continue
            
            # Detect headings (simple heuristic)
            if len(line) < 100 and (
                line.isupper() or
                re.match(r'^\d+\.?\s+[A-Z]', line) or
                re.match(r'^[A-Z][^.!?]*$', line)
            ):
                # Determine heading level based on context
                if any(keyword in line.lower() for keyword in ['section', 'chapter', 'part']):
                    markdown_lines.append(f'## {line}')
                else:
                    markdown_lines.append(f'### {line}')
            else:
                markdown_lines.append(line)
        
        return '\n'.join(markdown_lines)
    
    def _calculate_hash(self, content: str) -> str:
        """Calculate SHA-256 hash of content"""
        return hashlib.sha256(content.encode('utf-8')).hexdigest()
    
    def extract_structured_data(
        self,
        text: str,
        extraction_type: str = 'regulatory'
    ) -> Dict[str, Any]:
        """
        Extract structured data from text using NLP
        
        Args:
            text: Input text
            extraction_type: Type of extraction (regulatory, medical, general)
        
        Returns:
            Dictionary with extracted structured data
        """
        try:
            if not self.config.nlp_enabled:
                return {"success": False, "error": "NLP processing disabled"}
            
            nlp = get_nlp_model()
            doc = nlp(text)
            
            # Extract entities
            entities = []
            for ent in doc.ents:
                entities.append({
                    "text": ent.text,
                    "label": ent.label_,
                    "start": ent.start_char,
                    "end": ent.end_char,
                    "description": spacy.explain(ent.label_)
                })
            
            # Extract key phrases and concepts
            key_phrases = self._extract_key_phrases(text, extraction_type)
            
            # Extract regulatory-specific information
            regulatory_info = {}
            if extraction_type == 'regulatory':
                regulatory_info = self._extract_regulatory_concepts(text)
            
            result = ExtractionResult(
                entities=entities,
                key_phrases=key_phrases,
                regulatory_concepts=regulatory_info,
                confidence_score=self._calculate_extraction_confidence(entities, key_phrases)
            )
            
            return {
                "success": True,
                "extraction_result": asdict(result)
            }
            
        except Exception as e:
            logger.error(f"Error in structured data extraction: {str(e)}")
            return {
                "success": False,
                "error": str(e)
            }
    
    def _extract_key_phrases(self, text: str, extraction_type: str) -> List[str]:
        """Extract key phrases from text"""
        # Simple keyword extraction using TF-IDF
        # In production, this could use more sophisticated methods
        
        # Tokenize and clean text
        words = nltk.word_tokenize(text.lower())
        words = [word for word in words if word.isalpha() and len(word) > 3]
        
        # Remove stopwords
        from nltk.corpus import stopwords
        stop_words = set(stopwords.words('english'))
        words = [word for word in words if word not in stop_words]
        
        # Count frequency
        from collections import Counter
        word_freq = Counter(words)
        
        # Return top phrases
        return [word for word, count in word_freq.most_common(20)]
    
    def _extract_regulatory_concepts(self, text: str) -> Dict[str, Any]:
        """Extract FDA regulatory-specific concepts"""
        regulatory_patterns = {
            'k_numbers': re.findall(r'K\d{6}', text, re.IGNORECASE),
            'product_codes': re.findall(r'\b[A-Z]{3}\b', text),
            'cfr_sections': re.findall(r'21\s*CFR\s*\d+(?:\.\d+)?', text, re.IGNORECASE),
            'device_classes': re.findall(r'Class\s*[I]{1,3}', text, re.IGNORECASE),
            'fda_guidance': re.findall(r'FDA\s+Guidance', text, re.IGNORECASE),
            'predicate_devices': re.findall(r'predicate\s+device', text, re.IGNORECASE)
        }
        
        # Extract dates
        date_patterns = re.findall(
            r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b',
            text
        )
        regulatory_patterns['dates'] = date_patterns
        
        return regulatory_patterns
    
    def _calculate_extraction_confidence(self, entities: List[Dict], key_phrases: List[str]) -> float:
        """Calculate confidence score for extraction results"""
        # Simple confidence calculation based on number of entities and phrases found
        entity_score = min(len(entities) / 10, 1.0)  # Normalize to 0-1
        phrase_score = min(len(key_phrases) / 20, 1.0)  # Normalize to 0-1
        
        return (entity_score + phrase_score) / 2
    
    def search_documents(
        self,
        query: str,
        documents: List[Dict[str, Any]],
        top_k: int = 5
    ) -> Dict[str, Any]:
        """
        Search documents using semantic similarity
        
        Args:
            query: Search query
            documents: List of documents to search
            top_k: Number of top results to return
        
        Returns:
            Dictionary with search results
        """
        try:
            if not documents:
                return {
                    "success": True,
                    "results": [],
                    "total_documents": 0
                }
            
            # Extract text content from documents
            doc_texts = []
            for doc in documents:
                if isinstance(doc, dict):
                    text = doc.get('content', '') or doc.get('text', '') or str(doc)
                else:
                    text = str(doc)
                doc_texts.append(text)
            
            # Use sentence transformer for semantic search
            sentence_transformer = get_sentence_transformer()
            
            # Encode query and documents
            query_embedding = sentence_transformer.encode([query])
            doc_embeddings = sentence_transformer.encode(doc_texts)
            
            # Calculate similarities
            similarities = cosine_similarity(query_embedding, doc_embeddings)[0]
            
            # Get top results
            top_indices = np.argsort(similarities)[::-1][:top_k]
            
            results = []
            for idx in top_indices:
                if similarities[idx] > 0.1:  # Minimum similarity threshold
                    result = SearchResult(
                        document_id=documents[idx].get('id', str(idx)),
                        title=documents[idx].get('title', f'Document {idx}'),
                        content_preview=doc_texts[idx][:200] + '...' if len(doc_texts[idx]) > 200 else doc_texts[idx],
                        similarity_score=float(similarities[idx]),
                        metadata=documents[idx].get('metadata', {})
                    )
                    results.append(asdict(result))
            
            return {
                "success": True,
                "results": results,
                "total_documents": len(documents),
                "query": query
            }
            
        except Exception as e:
            logger.error(f"Error in document search: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "results": []
            }
    
    def summarize_document(
        self,
        text: str,
        max_length: int = 150,
        min_length: int = 50
    ) -> Dict[str, Any]:
        """
        Generate summary of document text
        
        Args:
            text: Input text to summarize
            max_length: Maximum summary length
            min_length: Minimum summary length
        
Returns:
            Dictionary with summary results
        """
        try:
            if len(text.strip()) < min_length:
                return {
                    "success": False,
                    "error": "Text too short to summarize"
                }
            
            # Split long text into chunks for processing
            max_chunk_length = 1024  # Model input limit
            chunks = self._split_text_into_chunks(text, max_chunk_length)
            
            summarizer = get_summarizer()
            summaries = []
            
            for chunk in chunks:
                if len(chunk.strip()) < min_length:
                    continue
                
                try:
                    summary = summarizer(
                        chunk,
                        max_length=max_length,
                        min_length=min_length,
                        do_sample=False
                    )
                    summaries.append(summary[0]['summary_text'])
                except Exception as e:
                    logger.warning(f"Error summarizing chunk: {str(e)}")
                    continue
            
            # Combine summaries if multiple chunks
            if len(summaries) > 1:
                combined_summary = ' '.join(summaries)
                # Summarize the combined summaries if too long
                if len(combined_summary) > max_length * 2:
                    final_summary = summarizer(
                        combined_summary,
                        max_length=max_length,
                        min_length=min_length,
                        do_sample=False
                    )
                    final_text = final_summary[0]['summary_text']
                else:
                    final_text = combined_summary
            else:
                final_text = summaries[0] if summaries else "Unable to generate summary"
            
            result = SummaryResult(
                summary=final_text,
                original_length=len(text),
                summary_length=len(final_text),
                compression_ratio=len(final_text) / len(text),
                confidence_score=0.8 if summaries else 0.0
            )
            
            return {
                "success": True,
                "summary_result": asdict(result)
            }
            
        except Exception as e:
            logger.error(f"Error in document summarization: {str(e)}")
            return {
                "success": False,
                "error": str(e)
            }
    
    def _split_text_into_chunks(self, text: str, max_length: int) -> List[str]:
        """Split text into chunks for processing"""
        # Split by sentences first
        sentences = nltk.sent_tokenize(text)
        
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            if len(current_chunk) + len(sentence) <= max_length:
                current_chunk += " " + sentence if current_chunk else sentence
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = sentence
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks
    
    def track_document_version(
        self,
        document_id: str,
        content: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Track document version and detect changes
        
        Args:
            document_id: Unique document identifier
            content: Document content
            metadata: Additional metadata
        
        Returns:
            Dictionary with version tracking results
        """
        try:
            content_hash = self._calculate_hash(content)
            
            # In a real implementation, this would interact with a database
            # For now, we'll simulate version tracking
            
            version = DocumentVersion(
                document_id=document_id,
                version_number=1,  # Would be incremented based on existing versions
                content_hash=content_hash,
                content_length=len(content),
                created_at=datetime.now(),
                metadata=metadata or {}
            )
            
            # Detect changes (simplified)
            changes_detected = True  # Would compare with previous version
            
            return {
                "success": True,
                "version": asdict(version),
                "changes_detected": changes_detected,
                "content_hash": content_hash
            }
            
        except Exception as e:
            logger.error(f"Error in version tracking: {str(e)}")
            return {
                "success": False,
                "error": str(e)
            }


# Utility functions for FDA-specific document processing

def extract_fda_guidance_metadata(text: str) -> Dict[str, Any]:
    """Extract metadata specific to FDA guidance documents"""
    metadata = {}
    
    # Extract document title
    title_match = re.search(r'^(.+?)(?:\n|$)', text.strip(), re.MULTILINE)
    if title_match:
        metadata['title'] = title_match.group(1).strip()
    
    # Extract effective date
    date_patterns = [
        r'Effective\s+Date[:\s]+([A-Za-z]+\s+\d{1,2},?\s+\d{4})',
        r'Date\s+of\s+this\s+Guidance[:\s]+([A-Za-z]+\s+\d{1,2},?\s+\d{4})',
        r'([A-Za-z]+\s+\d{1,2},?\s+\d{4})'
    ]
    
    for pattern in date_patterns:
        date_match = re.search(pattern, text, re.IGNORECASE)
        if date_match:
            metadata['effective_date'] = date_match.group(1)
            break
    
    # Extract document number
    doc_num_match = re.search(r'Document\s+Number[:\s]+(\S+)', text, re.IGNORECASE)
    if doc_num_match:
        metadata['document_number'] = doc_num_match.group(1)
    
    # Extract subject/topic
    subject_match = re.search(r'Subject[:\s]+(.+?)(?:\n|$)', text, re.IGNORECASE | re.MULTILINE)
    if subject_match:
        metadata['subject'] = subject_match.group(1).strip()
    
    return metadata


def identify_regulatory_sections(text: str) -> List[Dict[str, Any]]:
    """Identify and extract regulatory sections from FDA documents"""
    sections = []
    
    # Common FDA document section patterns
    section_patterns = [
        r'(I{1,3}\.?\s+[A-Z][^.!?]*(?:Background|Introduction|Purpose|Scope))',
        r'(\d+\.?\s+[A-Z][^.!?]*(?:Requirements|Recommendations|Guidance))',
        r'([A-Z]\.\s+[A-Z][^.!?]*)',
        r'(Section\s+\d+[^.!?]*)'
    ]
    
    for pattern in section_patterns:
        matches = re.finditer(pattern, text, re.MULTILINE | re.IGNORECASE)
        for match in matches:
            sections.append({
                'title': match.group(1).strip(),
                'start_position': match.start(),
                'type': 'section_header'
            })
    
    return sections


# Export main class and utility functions
__all__ = [
    'DocumentProcessingTool',
    'DocumentProcessingConfig',
    'extract_fda_guidance_metadata',
    'identify_regulatory_sections'
]